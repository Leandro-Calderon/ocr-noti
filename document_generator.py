from docx import Document


def generar_documento(
    palabras, numeros, ruta_salida="/home/lean/Desktop/nuevo_documento.docx"
):
    """Genera un documento de Word con las palabras y números clave."""
    doc = Document()
    doc.add_paragraph("Palabras clave:")
    for palabra in palabras:
        doc.add_paragraph(palabra)

    doc.add_paragraph("Números clave:")
    for numero in numeros:
        doc.add_paragraph(numero)

    doc.save(ruta_salida)
    return ruta_salida
